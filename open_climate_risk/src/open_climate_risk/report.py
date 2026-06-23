"""Word and Excel report generation."""
from pathlib import Path
import numpy as np

from .config import OUTPUTS_DIR, RETURN_PERIODS


# ── Excel ──────────────────────────────────────────────────────────────────────

def generate_excel(
    asset: dict,
    df_hazard,
    df_risk,
    ead: float,
    ead_pct: float,
    results: list,
    path: str = None,
) -> str:
    """Generate an Excel workbook with all risk screening results (5 sheets)."""
    from openpyxl import Workbook
    from openpyxl.styles import Font, PatternFill, Alignment, Border, Side
    from openpyxl.utils import get_column_letter

    path = path or str(OUTPUTS_DIR / 'flood_risk_data.xlsx')

    DARK_BLUE = '1F4E79'; MID_BLUE = '1565C0'; ALT_BLUE = 'EBF3FB'; WHITE = 'FFFFFF'

    def hfill(hex_): return PatternFill('solid', fgColor=hex_)
    def bside():     return Side(style='thin', color='CCCCCC')
    def bord():      b = bside(); return Border(left=b, right=b, top=b, bottom=b)

    def h(ws, row, col, text):
        c = ws.cell(row, col, text)
        c.font      = Font(name='Arial', bold=True, size=10, color=WHITE)
        c.fill      = hfill(DARK_BLUE)
        c.alignment = Alignment(horizontal='center', vertical='center', wrap_text=True)
        c.border    = bord()
        return c

    def d(ws, row, col, val, alt=False, align='left', bold=False, fmt=None):
        c = ws.cell(row, col, val)
        c.font      = Font(name='Arial', size=10, bold=bold)
        c.fill      = hfill(ALT_BLUE if alt else WHITE)
        c.alignment = Alignment(horizontal=align, vertical='center')
        c.border    = bord()
        if fmt: c.number_format = fmt
        return c

    def autofit(ws):
        for col in ws.columns:
            w = max((len(str(c.value or '')) for c in col), default=0)
            ws.column_dimensions[get_column_letter(col[0].column)].width = min(w + 3, 45)

    wb   = Workbook()
    rps  = RETURN_PERIODS
    bl   = results[0]['ead_median']

    # ── Sheet 1 : Résumé ──────────────────────────────────────────────────────
    ws = wb.active
    ws.title = 'Résumé'
    ws.sheet_properties.tabColor = DARK_BLUE
    ws.freeze_panes = 'A2'
    ws.merge_cells('A1:F1')
    ws['A1'] = f'Analyse de Risque Inondation Fluviale \u2014 {asset["name"]}'
    ws['A1'].font      = Font(name='Arial', bold=True, size=13, color=WHITE)
    ws['A1'].fill      = hfill(DARK_BLUE)
    ws['A1'].alignment = Alignment(horizontal='center', vertical='center')
    ws.row_dimensions[1].height = 26

    ws['A3'] = 'Informations sur l\'actif'
    ws['A3'].font = Font(name='Arial', bold=True, size=11, color=DARK_BLUE)
    ws.merge_cells('A3:B3')

    asset_rows = [
        ('Nom',                        asset['name']),
        ('Type',                       asset['asset_type'].title()),
        ('Valeur de remplacement (€)', asset['value_eur']),
        ('Surface plancher (m²)',      asset['floor_area_m2']),
        ('Valeur unitaire (€/m²)',     asset['value_eur'] / asset['floor_area_m2']),
        ('Longitude',                  asset['lon']),
        ('Latitude',                   asset['lat']),
    ]
    for i, (k, v) in enumerate(asset_rows):
        d(ws, 4+i, 1, k, alt=(i%2==0), bold=True)
        c = d(ws, 4+i, 2, v, alt=(i%2==0), align='right')
        if isinstance(v, (int, float)) and v > 100:
            c.number_format = '#,##0'

    r = 4 + len(asset_rows) + 1
    ws.cell(r, 1, 'Résultats EAD par scénario').font = Font(name='Arial', bold=True, size=11, color=DARK_BLUE)
    ws.merge_cells(f'A{r}:F{r}')
    r += 1
    for j, hdr in enumerate(['Scénario', 'Ensemble', 'EAD médiane (€/an)', 'Min (€/an)', 'Max (€/an)', 'Δ vs baseline']):
        h(ws, r, j+1, hdr)
    r += 1
    for i, res in enumerate(results):
        ensemble = 'WATCH' if len(res['models']) == 1 else f'Médiane ({len(res["models"])} GCMs)'
        delta    = '—' if i == 0 else f'{(res["ead_median"]-bl)/bl*100:+.1f}%'
        for j, (v, fmt, aln) in enumerate([
            (res['label'],      None,    'left'),
            (ensemble,          None,    'left'),
            (res['ead_median'], '#,##0', 'right'),
            (res['ead_min'],    '#,##0', 'right'),
            (res['ead_max'],    '#,##0', 'right'),
            (delta,             None,    'center'),
        ]):
            d(ws, r+i, j+1, v, alt=(i%2==0), align=aln, fmt=fmt)
    ws.column_dimensions['A'].width = 28
    ws.column_dimensions['B'].width = 24
    for col in 'CDEF': ws.column_dimensions[col].width = 18

    # ── Sheet 2 : Aléa Historique ─────────────────────────────────────────────
    ws2 = wb.create_sheet('Aléa_Historique')
    ws2.sheet_properties.tabColor = MID_BLUE
    ws2.freeze_panes = 'A2'
    for j, hdr in enumerate(['T (ans)', 'Prob. annuelle', 'Profondeur (m)']):
        h(ws2, 1, j+1, hdr)
    for i, (_, row) in enumerate(df_hazard.iterrows()):
        alt = i % 2 == 0
        d(ws2, 2+i, 1, int(row['return_period_yr']),   alt=alt, align='center')
        d(ws2, 2+i, 2, row['exceedance_prob_yr'],       alt=alt, align='center', fmt='0.0%')
        d(ws2, 2+i, 3, round(row['flood_depth_m'], 4), alt=alt, align='center', fmt='0.0000')
    autofit(ws2)

    # ── Sheet 3 : Risque Historique ───────────────────────────────────────────
    ws3 = wb.create_sheet('Risque_Historique')
    ws3.sheet_properties.tabColor = MID_BLUE
    ws3.freeze_panes = 'A2'
    for j, hdr in enumerate(['T (ans)', 'Prob. annuelle', 'Profondeur (m)', 'Fraction dommage', 'Dommage (€)']):
        h(ws3, 1, j+1, hdr)
    for i, (_, row) in enumerate(df_risk.iterrows()):
        alt = i % 2 == 0
        d(ws3, 2+i, 1, int(row['return_period_yr']),     alt=alt, align='center')
        d(ws3, 2+i, 2, row['exceedance_prob_yr'],         alt=alt, align='center', fmt='0.0%')
        d(ws3, 2+i, 3, round(row['flood_depth_m'], 4),   alt=alt, align='center', fmt='0.0000')
        d(ws3, 2+i, 4, round(row['damage_fraction'], 4), alt=alt, align='center', fmt='0.00%')
        d(ws3, 2+i, 5, round(row['damage_eur'], 0),      alt=alt, align='right',  fmt='#,##0')
    r_ead = 2 + len(df_risk)
    for j in range(5):
        c = ws3.cell(r_ead, j+1)
        c.fill = hfill(MID_BLUE); c.border = bord()
        c.font = Font(name='Arial', bold=True, size=10, color=WHITE)
    ws3.cell(r_ead, 1, 'EAD').alignment = Alignment(horizontal='center')
    ws3.cell(r_ead, 4, 'EAD (€/an)').alignment = Alignment(horizontal='right')
    c_ead = ws3.cell(r_ead, 5, round(ead, 0))
    c_ead.alignment = Alignment(horizontal='right')
    c_ead.number_format = '#,##0'
    autofit(ws3)

    # ── Sheet 4 : Scénarios Ensemble ──────────────────────────────────────────
    ws4 = wb.create_sheet('Scénarios_Ensemble')
    ws4.sheet_properties.tabColor = '2196F3'
    ws4.freeze_panes = 'A2'
    hdrs4 = ['Scénario', 'Année', 'T (ans)', 'Prob. annuelle',
             'Prof. médiane (m)', 'Prof. min (m)', 'Prof. max (m)',
             'Dommage médiane (€)', 'Dommage min (€)', 'Dommage max (€)']
    for j, hdr in enumerate(hdrs4): h(ws4, 1, j+1, hdr)
    row_idx = 2
    for res in results:
        for k, rp in enumerate(rps):
            alt = row_idx % 2 == 0
            for j, (v, fmt, aln) in enumerate([
                (res['label'],                         None,     'left'),
                (res['year'],                          '0',      'center'),
                (int(rp),                              '0',      'center'),
                (round(1/rp, 4),                       '0.0%',   'center'),
                (round(res['depths_median'][k], 4),    '0.0000', 'center'),
                (round(res['depths_min'][k], 4),       '0.0000', 'center'),
                (round(res['depths_max'][k], 4),       '0.0000', 'center'),
                (round(res['damages_median'][k], 0),   '#,##0',  'right'),
                (round(res['damages_min'][k], 0),      '#,##0',  'right'),
                (round(res['damages_max'][k], 0),      '#,##0',  'right'),
            ]):
                d(ws4, row_idx, j+1, v, alt=alt, align=aln, fmt=fmt)
            row_idx += 1
    row_idx += 1
    ws4.cell(row_idx, 1, 'EAD par scénario').font = Font(name='Arial', bold=True, size=11, color=DARK_BLUE)
    ws4.merge_cells(f'A{row_idx}:J{row_idx}')
    row_idx += 1
    for j, hdr in enumerate(['Scénario', 'EAD médiane (€/an)', 'Min (€/an)', 'Max (€/an)', 'Δ vs baseline']):
        h(ws4, row_idx, j+1, hdr)
    row_idx += 1
    for i, res in enumerate(results):
        delta = '—' if i == 0 else f'{(res["ead_median"]-bl)/bl*100:+.1f}%'
        for j, (v, fmt, aln) in enumerate([
            (res['label'],      None,    'left'),
            (res['ead_median'], '#,##0', 'right'),
            (res['ead_min'],    '#,##0', 'right'),
            (res['ead_max'],    '#,##0', 'right'),
            (delta,             None,    'center'),
        ]):
            d(ws4, row_idx+i, j+1, v, alt=(i%2==0), align=aln, fmt=fmt)
    autofit(ws4)

    # ── Sheet 5 : Données Brutes GCM ──────────────────────────────────────────
    ws5 = wb.create_sheet('Données_Brutes_GCM')
    ws5.sheet_properties.tabColor = '9C27B0'
    ws5.freeze_panes = 'A2'
    for j, hdr in enumerate(['Scénario', 'Modèle', 'Année', 'T (ans)', 'Prob. annuelle',
                              'Profondeur (m)', 'Fraction dommage', 'Dommage (€)', 'EAD (€/an)']):
        h(ws5, 1, j+1, hdr)
    row_idx = 2
    for res in results:
        for model in res['models']:
            depths_m  = res['depths_per_model'][model]
            damages_m = res['damages_per_model'][model]
            ead_m     = res['ead_per_model'][model]
            for k, rp in enumerate(rps):
                alt = row_idx % 2 == 0
                for j, (v, fmt, aln) in enumerate([
                    (res['label'],                          None,     'left'),
                    (model,                                 None,     'left'),
                    (res['year'],                           '0',      'center'),
                    (int(rp),                               '0',      'center'),
                    (round(1/rp, 4),                        '0.0%',   'center'),
                    (round(float(depths_m[k]), 4),          '0.0000', 'center'),
                    (round(float(damages_m[k]/asset['value_eur']), 4), '0.00%', 'center'),
                    (round(float(damages_m[k]), 0),         '#,##0',  'right'),
                    (round(float(ead_m), 0),                '#,##0',  'right'),
                ]):
                    d(ws5, row_idx, j+1, v, alt=alt, align=aln, fmt=fmt)
                row_idx += 1
    autofit(ws5)

    wb.save(path)
    print(f'Excel généré \u2192 {path}  ({len(wb.sheetnames)} onglets)')
    return path


# ── Word ───────────────────────────────────────────────────────────────────────

def generate_word(
    asset: dict,
    df_risk,
    ead: float,
    ead_pct: float,
    results: list,
    rp_map: int = 500,
    path: str = None,
    outputs_dir: Path = None,
) -> str:
    """Generate a professional Word report with all risk screening results."""
    from docx import Document as DocxDocument
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement

    path       = path or str(OUTPUTS_DIR / 'flood_risk_report.docx')
    img_dir    = Path(outputs_dir) if outputs_dir else OUTPUTS_DIR
    TODAY      = '14 mars 2026'
    DARK_BLUE  = RGBColor(0x1F, 0x4E, 0x79)
    MID_BLUE   = RGBColor(0x15, 0x65, 0xC0)
    GREY       = RGBColor(0x55, 0x55, 0x55)

    # ── helpers ────────────────────────────────────────────────────────────────
    def set_cell_bg(cell, hex_color):
        tcPr = cell._tc.get_or_add_tcPr()
        shd  = OxmlElement('w:shd')
        shd.set(qn('w:val'), 'clear'); shd.set(qn('w:color'), 'auto')
        shd.set(qn('w:fill'), hex_color)
        tcPr.append(shd)

    def set_para_border_bottom(para, color='1565C0', sz=4):
        pPr = para._p.get_or_add_pPr()
        pBdr = OxmlElement('w:pBdr')
        bot  = OxmlElement('w:bottom')
        bot.set(qn('w:val'), 'single'); bot.set(qn('w:sz'), str(sz))
        bot.set(qn('w:space'), '4');    bot.set(qn('w:color'), color)
        pBdr.append(bot); pPr.append(pBdr)

    def add_tab_stops(para, stops):
        pPr  = para._p.get_or_add_pPr()
        tabs = OxmlElement('w:tabs')
        for pos, val in stops:
            t = OxmlElement('w:tab')
            t.set(qn('w:val'), val); t.set(qn('w:pos'), str(pos))
            tabs.append(t)
        pPr.append(tabs)

    def add_page_number(para):
        for tag, instr in [('begin', None), (None, ' PAGE '), ('end', None)]:
            r = para.add_run()
            r.font.name = 'Arial'; r.font.size = Pt(8); r.font.color.rgb = GREY
            if tag:
                el = OxmlElement('w:fldChar'); el.set(qn('w:fldCharType'), tag)
            else:
                el = OxmlElement('w:instrText')
                el.set(qn('xml:space'), 'preserve'); el.text = instr
            r._r.append(el)

    def head_row(table, texts, widths_cm, dark=True):
        row = table.rows[0]
        for cell, text, w in zip(row.cells, texts, widths_cm):
            cell.width = Cm(w)
            set_cell_bg(cell, '1F4E79' if dark else 'D5E8F0')
            p = cell.paragraphs[0]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(text)
            run.bold = True; run.font.name = 'Arial'; run.font.size = Pt(10)
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF) if dark else DARK_BLUE

    def data_row(table, texts, aligns=None, bold_first=False, alt=False):
        row = table.add_row()
        aligns = aligns or ['left'] * len(texts)
        for i, (cell, text, align) in enumerate(zip(row.cells, texts, aligns)):
            if alt: set_cell_bg(cell, 'EBF3FB')
            p = cell.paragraphs[0]
            p.alignment = {'left': WD_ALIGN_PARAGRAPH.LEFT,
                           'center': WD_ALIGN_PARAGRAPH.CENTER,
                           'right': WD_ALIGN_PARAGRAPH.RIGHT}[align]
            r = p.add_run(str(text))
            r.font.name = 'Arial'; r.font.size = Pt(10)
            r.bold = bold_first and i == 0

    def add_heading(doc, text, level=1):
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(20 if level == 1 else 14)
        p.paragraph_format.space_after  = Pt(6)
        if level == 1:
            set_para_border_bottom(p)
        r = p.add_run(text)
        r.font.name  = 'Arial'
        r.font.size  = Pt(15 if level == 1 else 12)
        r.bold       = True
        r.font.color.rgb = MID_BLUE if level == 1 else DARK_BLUE

    def add_para(doc, text, italic=False, color=None, size=10):
        p = doc.add_paragraph()
        r = p.add_run(text)
        r.font.name = 'Arial'; r.font.size = Pt(size); r.italic = italic
        if color: r.font.color.rgb = color
        p.paragraph_format.space_before = Pt(2)
        p.paragraph_format.space_after  = Pt(4)

    def add_image(doc, img_path, width_cm, caption):
        if not Path(str(img_path)).exists():
            add_para(doc, f'[Image manquante : {img_path}]', italic=True, color=GREY)
            return
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        p.add_run().add_picture(str(img_path), width=Cm(width_cm))
        cap = doc.add_paragraph(caption)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
        cap.paragraph_format.space_after = Pt(12)
        for r in cap.runs:
            r.font.name = 'Arial'; r.font.size = Pt(9)
            r.italic = True; r.font.color.rgb = GREY

    # ── Document setup ─────────────────────────────────────────────────────────
    doc = DocxDocument()
    CW  = 9072  # content width in twips (16 cm)

    for section in doc.sections:
        section.page_width    = Cm(21); section.page_height   = Cm(29.7)
        section.left_margin   = section.right_margin  = Cm(2.5)
        section.top_margin    = Cm(2.0); section.bottom_margin = Cm(2.0)
        section.different_first_page_header_footer = True

        hpara = section.header.paragraphs[0]
        set_para_border_bottom(hpara, sz=4)
        add_tab_stops(hpara, [(CW, 'right')])
        for text in ['Analyse de Risque Inondation Fluviale', f'\t{TODAY}']:
            r = hpara.add_run(text)
            r.font.name = 'Arial'; r.font.size = Pt(8)
            r.italic = True; r.font.color.rgb = GREY

        fpara = section.footer.paragraphs[0]
        add_tab_stops(fpara, [(CW // 2, 'center'), (CW, 'right')])
        rf = fpara.add_run('Open Climate Risk')
        rf.font.name = 'Arial'; rf.font.size = Pt(8); rf.font.color.rgb = GREY
        rt = fpara.add_run('\t'); rt.font.name = 'Arial'; rt.font.size = Pt(8)
        add_page_number(fpara)
        rr = fpara.add_run(f'\t{asset["name"]}')
        rr.font.name = 'Arial'; rr.font.size = Pt(8); rr.font.color.rgb = GREY

    # ── Cover ──────────────────────────────────────────────────────────────────
    banner = doc.add_table(rows=2, cols=1)
    banner.style = 'Table Grid'
    c0 = banner.rows[0].cells[0]
    set_cell_bg(c0, '1F4E79')
    p0 = c0.paragraphs[0]; p0.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p0.paragraph_format.space_before = Pt(28); p0.paragraph_format.space_after = Pt(6)
    r = p0.add_run('ANALYSE DE RISQUE INONDATION FLUVIALE')
    r.font.name = 'Arial'; r.font.size = Pt(20); r.bold = True
    r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    c1 = banner.rows[1].cells[0]
    set_cell_bg(c1, '2E75B6')
    p1 = c1.paragraphs[0]; p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p1.paragraph_format.space_before = Pt(8); p1.paragraph_format.space_after = Pt(8)
    r2 = p1.add_run('Open Climate Risk  \u2014  Screening Notebook 01')
    r2.font.name = 'Arial'; r2.font.size = Pt(12)
    r2.font.color.rgb = RGBColor(0xDE, 0xEB, 0xF7)

    doc.add_paragraph()
    for text, size, bold, color in [
        (asset['name'], 22, True, DARK_BLUE),
        (f'{asset["asset_type"].title()}  \u00b7  {asset["floor_area_m2"]} m\u00b2  \u00b7  {asset["value_eur"]:,} \u20ac',
         13, False, GREY),
        (TODAY, 11, False, GREY),
    ]:
        p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(text)
        r.font.name = 'Arial'; r.font.size = Pt(size); r.bold = bold
        r.font.color.rgb = color

    doc.add_paragraph()
    toc_h = doc.add_paragraph()
    toc_h.paragraph_format.space_after = Pt(6)
    r = toc_h.add_run('Table des mati\u00e8res')
    r.font.name = 'Arial'; r.font.size = Pt(12); r.bold = True
    r.font.color.rgb = DARK_BLUE
    for num, title in [
        ('1.', 'M\u00e9thodologie'),
        ('2.', 'Actif analys\u00e9 & localisation'),
        ('3.', 'Cartes d\'al\u00e9a'),
        ('4.', 'Courbes de vuln\u00e9rabilit\u00e9 JRC'),
        ('5.', 'R\u00e9sultats historiques'),
        ('6.', 'Sc\u00e9narios climatiques 2050'),
        ('7.', 'Limites & R\u00e9f\u00e9rences'),
    ]:
        p = doc.add_paragraph()
        p.paragraph_format.space_before = Pt(2); p.paragraph_format.space_after = Pt(2)
        r = p.add_run(f'  {num}  {title}')
        r.font.name = 'Arial'; r.font.size = Pt(10); r.font.color.rgb = MID_BLUE

    doc.add_page_break()

    # ── 1. Méthodologie ───────────────────────────────────────────────────────
    add_heading(doc, '1. M\u00e9thodologie')
    add_para(doc, 'La m\u00e9thodologie combine trois composantes pour estimer l\'EAD :')
    add_para(doc, 'Al\u00e9a (WRI Aqueduct)  \u00d7  Vuln\u00e9rabilit\u00e9 (JRC 2017)  \u00d7  Exposition (actif)  \u2192  EAD',
             color=MID_BLUE)
    doc.add_paragraph()
    tbl = doc.add_table(rows=1, cols=4); tbl.style = 'Table Grid'
    head_row(tbl, ['Composante', 'Dataset', 'Licence', 'R\u00e9f\u00e9rence'], [3, 5.5, 3, 4.5])
    for i, row_data in enumerate([
        ('Al\u00e9a',          'WRI Aqueduct Floods v2',                  'CC BY 4.0',     'Ward et al. 2020'),
        ('Vuln\u00e9rabilit\u00e9', 'JRC Global Flood Depth-Damage Functions', 'Domaine public', 'Huizinga et al. 2017'),
        ('Exposition',    'D\u00e9fini par l\'utilisateur',               '\u2014',         '\u2014'),
    ]):
        data_row(tbl, row_data, bold_first=True, alt=(i % 2 == 0))
    doc.add_page_break()

    # ── 2. Actif & localisation ───────────────────────────────────────────────
    add_heading(doc, '2. Actif analys\u00e9 & localisation')
    tbl2 = doc.add_table(rows=1, cols=2); tbl2.style = 'Table Grid'
    head_row(tbl2, ['Param\u00e8tre', 'Valeur'], [7, 9])
    for i, (k, v) in enumerate([
        ('Nom',                    asset['name']),
        ('Type',                   asset['asset_type'].title()),
        ('Valeur de remplacement', f'{asset["value_eur"]:,} \u20ac'),
        ('Surface plancher',       f'{asset["floor_area_m2"]} m\u00b2'),
        ('Valeur unitaire',        f'{asset["value_eur"]/asset["floor_area_m2"]:,.0f} \u20ac/m\u00b2'),
        ('Coordonn\u00e9es',       f'{asset["lon"]:.4f}\u00b0E, {asset["lat"]:.4f}\u00b0N'),
    ]):
        data_row(tbl2, [k, str(v)], bold_first=True, alt=(i % 2 == 0))
    doc.add_paragraph()
    add_image(doc, img_dir / 'asset_location.png', 11,
              'Figure 1 \u2014 Localisation de l\'actif (\u00e9toile rouge). Fond\u202f: CartoDB Positron.')
    doc.add_page_break()

    # ── 3. Cartes d'aléa ─────────────────────────────────────────────────────
    add_heading(doc, f'3. Cartes d\'al\u00e9a (T\u202f=\u202f{rp_map} ans)')
    add_para(doc, f'Profondeur d\'inondation WRI Aqueduct Floods v2 \u2014 r\u00e9solution native ~1\u202fkm/pixel.')
    add_image(doc, img_dir / f'hazard_map_T{rp_map}_smooth.png', 14,
              f'Figure 2 \u2014 Carte d\'al\u00e9a T={rp_map} ans, interpolation bilin\u00e9aire (\u00d74, liss\u00e9e)')
    doc.add_paragraph()
    add_image(doc, img_dir / f'hazard_map_T{rp_map}_raw.png', 14,
              f'Figure 3 \u2014 Carte d\'al\u00e9a T={rp_map} ans, donn\u00e9es brutes (~1\u202fkm/pixel)')
    doc.add_page_break()

    # ── 4. JRC curves ─────────────────────────────────────────────────────────
    add_heading(doc, '4. Courbes de vuln\u00e9rabilit\u00e9 JRC')
    add_para(doc, f'Fonctions JRC Huizinga 2017 \u2014 Europe occidentale. '
             f'Courbe \u00ab\u202f{asset["asset_type"]}\u202f\u00bb mise en \u00e9vidence.')
    add_image(doc, img_dir / 'jrc_curves.png', 15,
              'Figure 4 \u2014 Fonctions de dommage JRC Huizinga 2017 (Europe occidentale)')
    doc.add_page_break()

    # ── 5. Résultats historiques ──────────────────────────────────────────────
    add_heading(doc, '5. R\u00e9sultats historiques par p\u00e9riode de retour')
    tbl3 = doc.add_table(rows=1, cols=5); tbl3.style = 'Table Grid'
    head_row(tbl3, ['T (ans)', 'Prob. ann.', 'Prof. (m)', 'Frac. dom.', 'Dommage (\u20ac)'],
             [2, 2.5, 2.5, 2.5, 4.5])
    for i, (_, row) in enumerate(df_risk.iterrows()):
        data_row(tbl3, [
            f'{int(row.return_period_yr)}',
            f'{row.exceedance_prob_yr:.1%}',
            f'{row.flood_depth_m:.2f}',
            f'{row.damage_fraction:.1%}',
            f'{row.damage_eur:,.0f}',
        ], aligns=['center', 'center', 'center', 'center', 'right'], alt=(i % 2 == 0))
    doc.add_paragraph()
    ead_p = doc.add_paragraph(); ead_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = ead_p.add_run(f'EAD\u202f=\u202f{ead:,.0f}\u202f\u20ac/an   |   EAD / Valeur actif\u202f=\u202f{ead_pct:.2f}\u202f%/an')
    r.font.name = 'Arial'; r.font.size = Pt(12); r.bold = True
    r.font.color.rgb = MID_BLUE
    doc.add_paragraph()
    add_image(doc, img_dir / 'flood_risk_screening.png', 15,
              'Figure 5 \u2014 Screening de risque inondation (baseline historique)')
    doc.add_page_break()

    # ── 6. Scénarios climatiques ──────────────────────────────────────────────
    add_heading(doc, '6. Sc\u00e9narios climatiques 2050')
    add_para(doc, '5 GCMs ISIMIP2b pour RCP\u202f4.5 et RCP\u202f8.5 \u00e0 l\'horizon 2050. '
             'Statistiques : m\u00e9diane, min et max inter-mod\u00e8les.')
    tbl4 = doc.add_table(rows=1, cols=5); tbl4.style = 'Table Grid'
    head_row(tbl4, ['Sc\u00e9nario', 'EAD m\u00e9diane (\u20ac/an)', 'Min (\u20ac/an)',
                    'Max (\u20ac/an)', '\u0394 vs baseline'], [5, 3.5, 3.5, 3.5, 3])
    bl_ead = results[0]['ead_median']
    for i, res in enumerate(results):
        delta = '\u2014' if i == 0 else f'{(res["ead_median"]-bl_ead)/bl_ead*100:+.1f}%'
        data_row(tbl4, [res['label'], f'{res["ead_median"]:,.0f}',
                        f'{res["ead_min"]:,.0f}', f'{res["ead_max"]:,.0f}', delta],
                 aligns=['left', 'right', 'right', 'right', 'center'], alt=(i % 2 == 0))
    doc.add_paragraph()
    add_para(doc, 'Note\u202f: La dispersion inter-mod\u00e8les domine sur la diff\u00e9rence entre '
             'RCP\u202f4.5 et RCP\u202f8.5 \u00e0 2050. Le risque de queue reste la m\u00e9trique cl\u00e9.',
             italic=True, color=GREY)
    doc.add_paragraph()
    add_image(doc, img_dir / 'scenarios_comparison.png', 15,
              'Figure 6 \u2014 Comparaison des sc\u00e9narios climatiques (m\u00e9diane GCM \u00b1 enveloppe min/max)')
    doc.add_page_break()
    add_image(doc, img_dir / 'hazard_curves_per_gcm.png', 15,
              'Figure 7 \u2014 Courbes d\'al\u00e9a brutes par GCM \u00d7 sc\u00e9nario (r\u00e9f. hist. WATCH en gris)')
    doc.add_page_break()

    # ── 7. Limites ────────────────────────────────────────────────────────────
    add_heading(doc, '7. Limites & R\u00e9f\u00e9rences')
    tbl5 = doc.add_table(rows=1, cols=3); tbl5.style = 'Table Grid'
    head_row(tbl5, ['Limitation', 'Impact', 'Mitigation'], [5, 5, 6])
    for i, row_data in enumerate([
        ('R\u00e9solution Aqueduct (~1\u202fkm)', 'Peut manquer les caract\u00e9ristiques locales',  'MNT + mod\u00e8le hydraulique local'),
        ('Courbes JRC = moyennes r\u00e9gionales', 'Peut ne pas refl\u00e9ter le b\u00e2ti local',  'Enqu\u00eates de dommages locales'),
        ('Analyse \u00e0 l\'\u00e9chelle d\'un actif', 'Pas de diversification portefeuille',        'Agr\u00e9ger au niveau portefeuille'),
        ('Horizon 2050 uniquement',             'Divergence sc\u00e9narios sous-estim\u00e9e',    'Ajouter horizons 2030 et 2080'),
    ]):
        data_row(tbl5, row_data, alt=(i % 2 == 0))
    doc.add_paragraph()
    add_heading(doc, 'R\u00e9f\u00e9rences', level=2)
    for ref in [
        'Ward, P.J. et al. (2020). Aqueduct Floods Methodology. Technical Note. WRI, Washington D.C.',
        'Huizinga, J., de Moel, H. & Szewczyk, W. (2017). Global flood depth-damage functions. EUR 28552 EN. JRC.',
    ]:
        add_para(doc, ref)

    doc.save(path)
    print(f'Word g\u00e9n\u00e9r\u00e9 \u2192 {path}')
    return path
